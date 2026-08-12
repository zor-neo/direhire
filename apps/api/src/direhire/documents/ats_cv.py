from __future__ import annotations

import io
from html import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

from direhire.ai.private_contracts import TailoredCvResult

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)


def _set_font(run: object, name: str, size: float, *, bold: bool = False) -> None:
    run.font.name = name  # type: ignore[attr-defined]
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)  # type: ignore[attr-defined]
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)  # type: ignore[attr-defined]
    run.font.size = Pt(size)  # type: ignore[attr-defined]
    run.bold = bold  # type: ignore[attr-defined]


class AtsCvRenderer:
    """Renders one restrained, parser-friendly P0 CV template."""

    def render(self, content: dict[str, object], format: str) -> bytes:
        result = TailoredCvResult.model_validate(content)
        if format == "DOCX":
            return self._docx(result)
        if format == "PDF":
            return self._pdf(result)
        raise ValueError("unsupported document format")

    def _docx(self, result: TailoredCvResult) -> bytes:
        document = Document()
        section = document.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.25
        for style_name, size, color, before, after in (
            ("Heading 1", 16, BLUE, 18, 10),
            ("Heading 2", 13, BLUE, 14, 7),
            ("Heading 3", 12, DARK_BLUE, 10, 5),
        ):
            style = document.styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        title = document.add_paragraph()
        title.paragraph_format.space_after = Pt(4)
        _set_font(title.add_run(result.title), "Calibri", 20, bold=True)
        summary_heading = document.add_paragraph("Professional Summary", style="Heading 1")
        summary_heading.paragraph_format.keep_with_next = True
        document.add_paragraph(result.professional_summary)
        for cv_section in result.sections:
            document.add_paragraph(cv_section.heading, style="Heading 1")
            for item in cv_section.items:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.left_indent = Inches(0.375)
                paragraph.paragraph_format.first_line_indent = Inches(-0.188)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.25
                paragraph.add_run(item)
        output = io.BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _pdf(result: TailoredCvResult) -> bytes:
        output = io.BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=LETTER,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
            title=result.title,
            author="",
            creator="DireHire",
        )
        base = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CvTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            alignment=TA_LEFT,
            textColor=HexColor("#000000"),
            spaceAfter=8,
        )
        heading_style = ParagraphStyle(
            "CvHeading",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=17,
            textColor=HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
            keepWithNext=True,
        )
        body_style = ParagraphStyle(
            "CvBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=11,
            leading=14,
            spaceAfter=6,
        )
        story: list[object] = [
            Paragraph(escape(result.title), title_style),
            Paragraph("Professional Summary", heading_style),
            Paragraph(escape(result.professional_summary), body_style),
        ]
        for cv_section in result.sections:
            story.append(Paragraph(escape(cv_section.heading), heading_style))
            story.append(
                ListFlowable(
                    [
                        ListItem(Paragraph(escape(item), body_style), leftIndent=0)
                        for item in cv_section.items
                    ],
                    bulletType="bullet",
                    start="circle",
                    leftIndent=27,
                    bulletFontName="Helvetica",
                    bulletFontSize=8,
                )
            )
            story.append(Spacer(1, 3))
        document.build(story)
        return output.getvalue()
