import io

import pytest
from direhire.ai.private_service import PrivateAiRequestService
from direhire.config import Settings
from direhire.documents.ats_cv import AtsCvRenderer
from direhire.documents.service import TailoredCvDocumentProcessor, TailoredCvService
from direhire.errors import NotFoundError
from direhire.files.storage import PresignedUpload, StoredObjectMetadata
from direhire.models import OutboxEvent, PrivateAiArtifact, PrivateFile, TailoredCvDocument, User
from docx import Document
from pypdf import PdfReader
from sqlalchemy import func, select
from sqlalchemy.orm import Session, sessionmaker

from tests.conftest import USER_A, USER_B

TAILORED = {
    "title": "Backend Engineer",
    "professional_summary": "Backend engineer experienced in Python API delivery.",
    "sections": [
        {
            "heading": "Experience",
            "items": [
                "Built documented Python APIs for a fictional internal platform.",
                "Improved PostgreSQL query reliability using measured production evidence.",
            ],
        },
        {"heading": "Skills", "items": ["Python", "PostgreSQL", "FastAPI"]},
    ],
    "omitted_or_deemphasized": [],
    "truthfulness_notes": ["All statements must remain grounded in the Base CV."],
}


class MemoryStorage:
    def __init__(self) -> None:
        self.objects: dict[str, tuple[bytes, str]] = {}
        self.downloads: list[str] = []

    def write(self, *, bucket: str, key: str, content: bytes, content_type: str) -> None:
        del bucket
        self.objects[key] = (content, content_type)

    def create_download(self, *, bucket: str, key: str, filename: str, expires_seconds: int) -> str:
        del bucket, expires_seconds
        self.downloads.append(f"{key}:{filename}")
        return "https://download.example.invalid/private"

    def create_upload(
        self, *, bucket: str, key: str, content_type: str, max_bytes: int
    ) -> PresignedUpload:
        raise NotImplementedError

    def head(self, *, bucket: str, key: str) -> StoredObjectMetadata:
        raise NotImplementedError

    def read(self, *, bucket: str, key: str, max_bytes: int) -> bytes:
        raise NotImplementedError

    def delete(self, *, bucket: str, key: str) -> None:
        del bucket
        self.objects.pop(key, None)

    def promote(self, *, bucket: str, source_key: str, destination_key: str) -> None:
        raise NotImplementedError


def seed_tailored(session: Session) -> str:
    session.add_all(
        [
            User(
                id=str(USER_A),
                cognito_subject="tailored-owner",
                email="tailored@example.invalid",
                plan="PREMIUM",
            ),
            User(
                id=str(USER_B),
                cognito_subject="tailored-other",
                email="other@example.invalid",
            ),
        ]
    )
    artifact = PrivateAiArtifact(
        user_id=str(USER_A),
        artifact_type="TAILORED_CV",
        idempotency_key="tailored-document-source",
        status="SUCCEEDED",
        input_hash="a" * 64,
        input_snapshot={"cv_text": "private source"},
        content=TAILORED,
        working_draft=TAILORED,
        name="Backend CV",
    )
    session.add(artifact)
    session.commit()
    return artifact.id


def test_ats_renderer_creates_parseable_docx_and_pdf() -> None:
    renderer = AtsCvRenderer()
    docx_bytes = renderer.render(TAILORED, "DOCX")
    pdf_bytes = renderer.render(TAILORED, "PDF")

    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(io.BytesIO(pdf_bytes)).pages
    )

    assert docx_bytes.startswith(b"PK")
    assert pdf_bytes.startswith(b"%PDF")
    assert "Backend Engineer" in text and "PostgreSQL" in text
    assert "Backend Engineer" in pdf_text and "PostgreSQL" in pdf_text
    section = document.sections[0]
    assert round(section.left_margin.inches, 2) == 1.0
    assert document.styles["Normal"].font.name == "Calibri"
    assert document.styles["Normal"].font.size.pt == 11
    assert document.styles["Heading 1"].font.size.pt == 16
    assert str(document.styles["Heading 1"].font.color.rgb) == "2E74B5"
    assert not document.tables


def test_versioned_documents_are_coalesced_private_and_downloadable(
    session_factory: sessionmaker[Session],
) -> None:
    storage = MemoryStorage()
    settings = Settings(environment="test", private_bucket_name="private")
    with session_factory() as database:
        artifact_id = seed_tailored(database)
        service = TailoredCvService(database)
        clone = service.duplicate(artifact_id, str(USER_A), "Backend CV v2")
        assert clone.version_number == 2
        assert clone.parent_artifact_id == artifact_id

        requested = service.request_document(clone.id, str(USER_A), "DOCX", correlation_id="d" * 36)
        repeated = service.request_document(clone.id, str(USER_A), "DOCX", correlation_id="e" * 36)
        assert repeated.id == requested.id
        assert database.scalar(select(func.count()).select_from(OutboxEvent)) == 1

        completed = TailoredCvDocumentProcessor(database, storage, settings).process(requested.id)
        assert completed.status == "SUCCEEDED"
        assert len(storage.objects) == 1
        private_file = database.get(PrivateFile, completed.file_id)
        assert private_file is not None and private_file.status == "CLEAN"
        assert service.download(completed.id, str(USER_A), storage, settings).startswith("https://")
        with pytest.raises(NotFoundError):
            service.download(completed.id, str(USER_B), storage, settings)
        assert database.scalar(select(func.count()).select_from(TailoredCvDocument)) == 1
        PrivateAiRequestService(database).delete(clone.id, str(USER_A), storage)
        assert storage.objects == {}
        assert database.scalar(select(func.count()).select_from(TailoredCvDocument)) == 0
        assert database.scalar(select(func.count()).select_from(PrivateFile)) == 0
